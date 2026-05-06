from docx import Document

## required python-docx==1.2.0

def extract_docx_content(file_path:str)->str:
    """
    Read the document with docx extention

    Args:
    - file_path (str): The path to the docx file.

    Returns:
    - str: The extracted text from the image.
    """
    try:
        # Load the document
        doc = Document(file_path)
        
        formatted_content = ""
    
        paragraph_index = 0
        table_index = 0
    
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 'p' stands for paragraph
                paragraph = doc.paragraphs[paragraph_index]
                formatted_content += paragraph.text + "\n\n"  # Add paragraph text with extra newlines
                paragraph_index += 1
            
            elif element.tag.endswith('tbl'):  # 'tbl' stands for table
                table = doc.tables[table_index]
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]  # Get text from each cell
                    formatted_content += " | ".join(row_data) + "\n"  # Join cells with " | "
                formatted_content += "\n"  # Add extra newline after each table
                table_index += 1
    
        # Return the formatted content
        return formatted_content
    except Exception as e:
        return f"Error occurred: {str(e)}"

# print(read_docx(r"../AI Invoice Auditor Agentic AI Capstone Project.docx"))
